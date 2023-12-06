from tkinter import *
from tkinter import ttk
from PIL import ImageTk, Image
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm
from openpyxl.utils import get_column_letter

# ---------------------------------
# Xử lý dữ liệu
# ---------------------------------
import pandas as pd
from numpy import array
import matplotlib.pyplot as plt
import numpy as np


def analysic():
    df = pd.read_csv('Student_Performance.csv', index_col=0, header=0)

    in_data = array(df.iloc[:, :])
    print(in_data)
    print('Tong so lop :')
    solop = in_data[:, 0]
    print(len(solop))
    print('Tong so sinh vien di thi :')
    tongsv = in_data[:, 1]
    print(np.sum(tongsv))
    diemA = in_data[:, 3]
    diemBc = in_data[:, 4]
    print('Tong sv:', tongsv)
    print("bang thong ke moi")
    in_datanew = np.delete(in_data, 2, axis=1)
    print(in_datanew)
    kieusosanh = ["A", "B+", "B", "C+", "C", "D+", "D", "F"]
    dulieusosanh = np.array([np.sum(in_data[:, 3]), np.sum(in_data[:, 4]), np.sum(in_data[:, 5])
                                , np.sum(in_data[:, 6]), np.sum(in_data[:, 7]), np.sum(in_data[:, 8])
                                , np.sum(in_data[:, 9]), np.sum(in_data[:, 10])])
    myexplode = [0.2, 0, 0, 0, 0, 0, 0, 0.2]
    fig1 = plt.figure()
    plt.pie(dulieusosanh, labels=kieusosanh, wedgeprops={'edgecolor': 'white', 'linewidth': 1.5}, explode=myexplode,
            autopct='%1.1f%%')
    plt.title("Tỉ lệ điểm của tất cả các lớp")
    plt.legend(loc='upper left', bbox_to_anchor=(0.9, 1.1))
    plt.axis('equal')
    plt.savefig('bieu_do_phan_tram.png')

    maxa = diemA.max()
    i, = np.where(diemA == maxa)
    print("=========")
    print('lop co nhieu diem A la {0} co {1} sv dat diem A'.format(in_data[i, 0], maxa))

    dict1 = dict(zip(kieusosanh, dulieusosanh))
    timmax = max(dict1.values())
    maxkey = [key for key, value in dict1.items() if value == timmax][0]
    print(f"Số điểm sinh viên đạt được nhiều nhất: {maxkey}")

    fig2 = plt.figure()
    kieusosanh1 = ["L1", "L2", "TX1", "TX2", "CK"]
    Solieu = np.array([np.sum(in_data[:, 11]), np.sum(in_data[:, 12]), np.sum(in_data[:, 13]), np.sum(in_data[:, 14]),
                       np.sum(in_data[:, 15])])
    plt.bar(kieusosanh1, Solieu, width=0.3)
    plt.title("Tổng các cơ số điểm")
    plt.xlabel('Loại bài thi')
    plt.ylabel(' Số sinh viên ')
    plt.savefig('bieu_do_cot.png')

    fig3 = plt.figure()
    plt.plot(range(len(diemA)), diemA, 'r-', label="Diem A")
    plt.plot(range(len(diemBc)), diemBc, 'g-', label="Diem B +")
    plt.xlabel('Lơp')
    plt.ylabel(' So sv dat diem ')
    plt.legend(loc='upper right')
    plt.savefig('bieu_do_duong.png')
    # plt.show()
    return


# ---------------------------------
# ---------------------------------


# ---------------------------------
# Giao diện
# ---------------------------------
class App:
    def __init__(self):
        self.root = Tk()
        self.root.title('Final Report')
        self.root.geometry('1000x700')

        # header text style
        self.header_font = ("Times New Roman", 20, 'bold')
        self.text_font = ("Times New Roman", 14, 'normal')

        # biểu đồ, tỉ lệ ảnh là 4:3
        self.score_ratio = ImageTk.PhotoImage(Image.open('score_ratio_1.png').resize((400, 300)))
        self.test_ratio = ImageTk.PhotoImage(Image.open('test_ratio.png').resize((400, 300)))

        self.style = ttk.Style()
        self.style.configure("TFrame", background="#fff")

        self.style1 = ttk.Style()
        self.style1.configure("Frame2.TFrame", background="#f0f0f0")

    # ---------------------------------
    # ---------------------------------
    # frame chứa định dạng báo cáo
    def draw_ui(self):
        # Tạo Frame chứa nội dung văn vản và hình ảnh với thanh cuộn
        self.top_frame = Frame(self.root, width=950, height=600, borderwidth=1)
        self.top_frame.pack(side=TOP)

        # Canvas
        self.canvas = Canvas(self.top_frame, background='#fff', width=900, height=600)
        self.canvas.pack(side=LEFT)

        # Scorllbar
        self.y_scrollbar = ttk.Scrollbar(self.top_frame, orient=VERTICAL, command=self.canvas.yview)
        self.y_scrollbar.pack(side=RIGHT, fill=Y)

        self.canvas.configure(yscrollcommand=self.y_scrollbar.set)

        # Content Frame
        self.content_frame = ttk.Frame(self.canvas, width=900, height=1600, style="TFrame")
        self.canvas.create_window((0, 0), window=self.content_frame, anchor=NW)

        self.header_text = Label(self.content_frame, text='Báo cáo môn học Lập trình Python (FE6051)', anchor=CENTER,
                                 justify=CENTER, background='#fff')
        self.header_text.configure(font=self.header_font)
        self.header_text.place(x=450, y=20, anchor=CENTER)

        self.body_text = Label(self.content_frame, text='I. Thông tin chung', background='#fff', font=self.text_font,
                               anchor=CENTER, justify=CENTER)
        self.body_text.place(x=50, y=100)

        self.body_text1 = Label(self.content_frame, text='Tên Môn: Lập trình Python', background='#fff',
                                font=self.text_font, anchor=CENTER, justify=CENTER)
        self.body_text1.place(x=50, y=130)

        self.body_text2 = Label(self.content_frame, text='Mã môn: FE6051        Số tín chỉ: 3 (2,1,0)',
                                background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        self.body_text2.place(x=50, y=160)

        self.body_text3 = Label(self.content_frame, text='Số lớp học phần: 9', background='#fff', font=self.text_font,
                                anchor=CENTER, justify=CENTER)
        self.body_text3.place(x=50, y=190)

        self.body_text4 = Label(self.content_frame, text='Tổng số sinh viên: 700        Pass: 83.2%', background='#fff',
                                font=self.text_font, anchor=CENTER, justify=CENTER)
        self.body_text4.place(x=50, y=220)

        self.body_text5 = Label(self.content_frame, text='II. Kết quả xử lý số liệu', background='#fff',
                                font=self.text_font, anchor=CENTER, justify=CENTER)
        self.body_text5.place(x=50, y=250)

        self.image1 = Label(self.content_frame, image=self.score_ratio, background='#fff', anchor=CENTER,
                            justify=CENTER)
        self.image1.place(x=450, y=450, anchor=CENTER)

        self.image2 = Label(self.content_frame, image=self.test_ratio, background='#fff', anchor=CENTER, justify=CENTER)
        self.image2.place(x=450, y=800, anchor=CENTER)

        self.body_text6 = Label(self.content_frame, text='III. Kết luận', background='#fff', font=self.text_font,
                                anchor=CENTER, justify=CENTER)
        self.body_text6.place(x=50, y=1000)

        self.body_text7 = Label(self.content_frame, text='.' * 79, background='#fff', font=self.text_font,
                                anchor=CENTER, justify=CENTER)
        self.body_text7.place(x=50, y=1030)

        self.body_text8 = Label(self.content_frame, text='.' * 79, background='#fff', font=self.text_font,
                                anchor=CENTER, justify=CENTER)
        self.body_text8.place(x=50, y=1060)

        # end report
        self.body_text9 = Label(self.content_frame, text='Ngày .... Tháng .... Năm .... ', background='#fff',
                                font=self.text_font, anchor=CENTER, justify=CENTER)
        self.body_text9.place(x=500, y=1100)

        self.body_text9 = Label(self.content_frame, text='Ký tên', background='#fff', font=self.text_font,
                                anchor=CENTER, justify=CENTER)
        self.body_text9.place(x=575, y=1130)

        self.content_frame.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

        # Tạo Frame chứa các Button điều khiển
        self.control_frame = ttk.Frame(self.root, style='Frame2.TFrame')
        self.control_frame.pack(fill=BOTH)

        self.word_btn = ttk.Button(self.control_frame, text='Export Word', command=self.export_word)
        self.word_btn.pack(side=LEFT, padx=10, pady=10, ipadx=10, ipady=5)

        self.excel_btn = ttk.Button(self.control_frame, text='Export Excel', command=self.export_excel)
        self.excel_btn.pack(side=LEFT, padx=10, pady=10, ipadx=10, ipady=5)

    # ---------------------------------
    # ---------------------------------

    # ---------------------------------

    # ---------------------------------
    def export_word(self):
        # Tạo một tài liệu mới
        doc = Document()

        # Đặt kích thước giấy A4
        section = doc.sections[0]
        section.page_width = Cm(21)  # Chiều rộng A4 (21 cm)
        section.page_height = Cm(29.7)  # height 29.7 cm

        # Định dạng tiêu đề
        title = doc.add_paragraph("Báo cáo môn học Lập trình Python (FE6051)")
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(18)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        body_text = [
            "I.Thông tin chung", "Tên môn: Lập trình Python",
            "Mã môn: FE6051        Số tín chỉ: 3(2,1,0)",
            "Số lớp học phần: 9",
            "Tổng số sinh viên: 700        Pass: 83.2%",
            "II. Kết quả xử lý số liệu"
        ]

        for i in body_text:
            body_para = doc.add_paragraph(i)
            body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            body_para.paragraph_format.line_spacing = Pt(15)  # dãn dòng 1.5
            body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            body_para.runs[0].font.size = Pt(14)

        # Chèn hình ảnh
        doc.add_picture('./score_ratio_1.png')
        doc.add_picture('./test_ratio.png')

        end_text = [
            "III. Kết luận",
            "." * 79,
            "." * 79
        ]

        for i in end_text:
            end_para = doc.add_paragraph(i)
            end_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            end_para.paragraph_format.line_spacing = Pt(15)  # dãn dòng 1.5
            end_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            end_para.runs[0].font.size = Pt(14)

        footer = "Ngày .... Tháng .... Năm .... \n\nKý tên\t\t"
        footer_doc = doc.add_paragraph(footer)
        footer_doc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_doc.paragraph_format.line_spacing = Pt(15)  # dãn dòng 1.5
        footer_doc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        footer_doc.runs[0].font.size = Pt(14)

        # Đặt font family cho toàn tài liệu
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"

        # Lưu tài liệu
        doc.save("Final_report.docx")

    # ---------------------------------

    # ---------------------------------


    # ---------------------------------
    def export_excel(self):
        # Workbook là nơi chứa tất cả các phần của tài liệu excel
        wb = Workbook()

        # Chọn trang tính mặc định
        sheet = wb.active

        # Đặt tiêu đề cho báo cáo
        sheet.title = "Lập trình Python (FE6051)"

        # Tạo font cho tiêu đề
        title_font = Font(name='Arial', size=16, bold=True)

        sheet.append(["Báo cáo học phần Lập trình Python (FE6051)", ])
        sheet.row_dimensions[1].height = 30  # chiều cao dòng 1 là 30px
        # merge dòng 1
        sheet.merge_cells('A1:R1')

        # căn giữa row đầu tiên
        for row in sheet.iter_rows(min_row=1, max_row=1, min_col=1, max_col=1):
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center')

        # Từ A đến R là định dạng chiều rộng A4
        # 50 dòng là chiều dài A4

        # Thêm dữ liệu vào báo cáo
        data = [
            ["Tổng số sinh viên", "Số sinh viên đạt A", "Số sinh viên điểm F", "Tỉ lệ sinh viên vượt qua học phần"],
            [519, "86 (16.6%)", "87 (16.8%)", "432 (83.2%)"],
        ]

        total_cells = 18
        num_columns = 4

        # Ghi dữ liệu vào các ô
        for row_data in data:
            sheet.append(row_data)

        # Định dạng dòng đầu tiên (header) với màu nền và in đậm
        # các dòng trong sheet đánh số từ 1 (khác với array đánh số từ 0)
        header_row = sheet[1]
        for cell in header_row:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="f0f0f0", end_color="f0f0f0", fill_type='solid')

        # Lưu tệp excel với tên file
        wb.save('final_report.xlsx')

    # ---------------------------------
    # ---------------------------------

    # làm cho GUI khi hiển thị sẽ căn giữa màn hình
    def make_center(self):
        self.root.update_idletasks()
        # lấy chiều rộng và cao của window tkinter
        width = self.root.winfo_width()
        height = self.root.winfo_height()

        # lấy chiều rộng màn hình laptop chia 2 trừ đi chiều rộng cửa sổ tkinter / 2
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry("{}x{}+{}+{}".format(width, height, x, y))


# ---------------------------------
# ---------------------------------

if __name__ == '__main__':
    app = App()
    app.make_center()
    app.draw_ui()
    app.root.mainloop()
